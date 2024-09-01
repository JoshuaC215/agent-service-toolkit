import aiohttp
import aiosqlite
import sqlite3
import json
import os
from docx import Document
from typing import AsyncGenerator, Dict, Any, Generator
import requests
from schema import ChatMessage, UserInput, StreamInput, Feedback

from api.wechat_rpa_api import WxAddGroupMember, WxDeleGroupMember, WxSetGroupNotice, DeviceData

# self.project_info群公告
# self.project_id查找群成员和初始化“进度表”和“问题表”
# self.status记录阶段
# self.customer_group、self.designer_group记录
class ConsultantAgentClient:
    def __init__(self, project_name: str, rpa_url: str = "http://101.43.11.104:8000/api/job/add_job", device_code: str = "2ea09832", base_url: str = "http://localhost:80"):
        """
        初始化客户端。

        Args:
            project_name (str): 项目名称。
            rpa_url (str): RPA的URL。
            device_code (str): 设备号。
            base_url (str): 智能体服务的基本 URL。
        """
        self.base_url = base_url    # 设置基础 URL
        self.rpa_url = rpa_url
        self.device_code = device_code
        self.auth_secret = os.getenv("AUTH_SECRET") # 从环境变量中获取认证密钥
        self.project_name = project_name
        self.db_path = "checkpoints.db"
        self.project_id, self.status, self.notice = self.fetch_project()
        if self.status == "初始化":
            self.initialize_project()

    @property
    def _headers(self):
        headers = {}    # 初始化请求头
        if self.auth_secret:    # 如果存在认证密钥
            headers["Authorization"] = f"Bearer {self.auth_secret}" # 添加认证头
        return headers  # 返回请求头

    # 读取数据库信息
    # project_info项目工单信息，初始化时编写群公告
    # completion_status项目状态信息
    # 初始化“进度表”和“问题表”
    # customer_group、designer_group检查成员变化
    def fetch_project(self):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT project_id, completion_status, notice FROM projects WHERE project_name = ?
                    """,
                    (self.project_name,),
                )
                row = cursor.fetchone()
                if row:
                    project_id, completion_status, notice = row
                    return project_id, completion_status, notice
                else:
                    print("没有找到该项目的信息")
                    return None, None, None
        except sqlite3.Error as e:
            print(f"数据库查询错误：{e}")
            return None, None, None

    # 初始化项目
    def initialize_project(self):
        # 操作微信群聊设置公告：获取工单信息、设置公告
        # 操作设置人员变动：获取群人员信息、人员变动
        self.notice = self.get_notice()
        if self.notice:
            self.update_group_notices(self.notice)

        self.update_project_info(completion_status="进行中")
        self.status = "进行中"

    def get_notice(self):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT project_type, project_period, video_duration, software_requirements,
                        reference_link, animation_text_content, voiceover_content,
                        project_name, completion_status
                    FROM projects WHERE project_id = ?
                    """,
                    (self.project_id,),
                )
                row = cursor.fetchone()
                if row:
                    project_info = dict(zip([
                        "project_type", "project_period", "video_duration", "software_requirements",
                        "reference_link", "animation_text_content", "voiceover_content",
                        "project_name", "completion_status"
                    ], row))

                    fields = [
                        ("project_type", "项目类型"),
                        ("project_period", "项目周期"),
                        ("video_duration", "视频时长"),
                        ("software_requirements", "软件要求"),
                        ("reference_link", "参考链接"),
                        ("animation_text_content", "动画文本内容文档"),
                        ("voiceover_content", "配音内容"),
                    ]

                    notice_content = "项目信息：\n\n"
                    for field, chinese_name in fields:
                        if project_info.get(field):
                            notice_content += f"{chinese_name}: {project_info[field]}\n"

                    return notice_content.strip()
                else:
                    print(f"未找到项目 ID {self.project_id} 的信息")
                    return None
        except sqlite3.Error as e:
            print(f"数据库查询错误：{e}")
            return None

    def update_group_notices(self, notice_content):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT customer_group, designer_group FROM
                    (SELECT DISTINCT customer_group FROM customers_group WHERE project_id = ?) AS c,
                    (SELECT DISTINCT designer_group FROM designers_group WHERE project_id = ?) AS d
                    """,
                    (self.project_id, self.project_id),
                )
                row = cursor.fetchone()
                if row:
                    customer_group, designer_group = row
                    self.set_group_notice(notice_content, customer_group)
                    self.set_group_notice(notice_content, designer_group)
                    print(f"已更新客户群 '{customer_group}' 和设计师群 '{designer_group}' 的公告")
                else:
                    print("未找到客户群或设计师群信息")
        except sqlite3.Error as e:
            print(f"获取群信息时发生错误：{e}")

    def update_group_members(self):
        customer_group, customer_members = self.get_group_info('customers_group')
        designer_group, designer_members = self.get_group_info('designers_group')
        if customer_group:
            self.sync_group_members(customer_group, customer_members)
        if designer_group:
            self.sync_group_members(designer_group, designer_members)

    def get_group_info(self, table_name):
        """
        获取群人员信息
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    f"""
                    SELECT {table_name.rstrip('s')}_group, {table_name.rstrip('s')}_id, {table_name.rstrip('s')}_name
                    FROM {table_name} WHERE project_id = ?
                    """,
                    (self.project_id,),
                )
                rows = cursor.fetchall()
                if rows:
                    group_name = rows[0][0]
                    members = [(row[1], row[2]) for row in rows]
                    return group_name, members
                return None, []
        except sqlite3.Error as e:
            print(f"获取{table_name}信息时发生错误：{e}")
            return None, []


# -----------------------------RPA request-----------------------------


    def set_group_notice(self, notice_content, group_name):
        """
        微信群聊设置公告
        """
        wx_set_group_notice = WxSetGroupNotice(
            params={
                "groupName": group_name,
                "content": notice_content
            }
        )
        device_data = DeviceData(
            device_code=self.device_code,
            job_list=[wx_set_group_notice]
        )
        try:
            response_text = device_data.send_request(self.rpa_url)
            print(f"群 '{group_name}' 公告设置成功, 响应:", response_text)
        except requests.RequestException as e:
            print(f"设置群 '{group_name}' 公告时发生错误：{e}")

    def add_group_members(self, group_name, members):
        """
        微信群聊添加多位成员
        """
        for member in members:
            wx_add_group_member = WxAddGroupMember(
                params={
                    "groupName": group_name,    # 群聊昵称
                    "phone": member['id'],      # 微信号
                }
            )
            device_data = DeviceData(
                device_code=self.device_code,
                job_list=[wx_add_group_member]
            )
            try:
                response_text = device_data.send_request(self.rpa_url)
                print(f"添加成员 {member['name']} 成功，响应：", response_text)
            except requests.RequestException as e:
                print(f"添加成员 {member['name']} 失败：{e}")

    def remove_group_members(self, group_name, members):
        """
        微信群聊删除多位成员
        """
        for member in members:
            wx_dele_group_member = WxDeleGroupMember(
                params={
                    "groupName": group_name,
                    "wxAccount": member['id'],
                }
            )
            device_data = DeviceData(
                device_code=self.device_code,
                job_list=[wx_dele_group_member]
            )
            try:
                response_text = device_data.send_request(self.rpa_url)
                print(f"删除成员 {member['name']} 成功，响应：", response_text)
            except requests.RequestException as e:
                print(f"删除成员 {member['name']} 失败：{e}")




    def update_project_info(self, **kwargs):
        set_clause = ", ".join([f"{key} = ?" for key in kwargs])
        values = list(kwargs.values())
        values.append(self.project_id)

        sql_query = f"UPDATE projects SET {set_clause} WHERE project_id = ?"

        try:
            with sqlite3.connect("checkpoints.db") as conn:
                cursor = conn.cursor()
                cursor.execute(sql_query, values)
                conn.commit()
                print("项目信息已更新。")
        except sqlite3.Error as e:
            print(f"更新项目信息时发送错误：{e}")
























    async def ainvoke(self, message: str, model: str|None = None, thread_id: str|None = None) -> ChatMessage:
        """
        异步调用代理。只返回最终消息。

        Args:
            message (str): 要发送给代理的消息。
            model (str, optional): 要使用的 LLM 模型。
            thread_id (str, optional): 继续对话的线程 ID。

        Returns:
            AnyMessage: 代理的响应。
        """
        async with aiohttp.ClientSession() as session:  # 创建异步 HTTP 会话
            request = UserInput(message=message)    # 创建请求对象
            if thread_id:   # 如果提供了线程 ID
                request.thread_id = thread_id   # 设置线程 ID
            if model:   # 如果提供了模型
                request.model = model   # 设置模型
            async with session.post(f"{self.base_url}/invoke", json=request.dict(), headers=self._headers) as response:
                if response.status == 200:  # 如果请求成功
                    result = await response.json()  # 解析 JSON 响应
                    return ChatMessage.parse_obj(result)    # 返回解析后的消息
                else:
                    raise Exception(f"Error: {response.status} - {await response.text()}")  # 抛出异常

    def invoke(self, message: str, model: str|None = None, thread_id: str|None = None) -> ChatMessage:
        """
        同步调用代理。只返回最终消息。

        Args:
            message (str): 要发送给代理的消息。
            model (str, optional): 要使用的 LLM 模型。
            thread_id (str, optional): 继续对话的线程 ID。

        Returns:
            ChatMessage: 代理的响应。
        """
        request = UserInput(message=message)    # 创建请求对象
        if thread_id:   # 如果提供了线程 ID
            request.thread_id = thread_id   # 设置线程 ID
        if model:   # 如果提供了模型
            request.model = model   # 设置模型
        response = requests.post(f"{self.base_url}/invoke", json=request.dict(), headers=self._headers) # 发送 POST 请求
        if response.status_code == 200: # 如果请求成功
            return ChatMessage.parse_obj(response.json())   # 返回解析后的消息
        else:
            raise Exception(f"Error: {response.status_code} - {response.text}") # 抛出异常


    # 创建项目文档
    def create_project_document(self):
        try:
            doc = Document()
            doc.add_heading("项目工单信息", level=1)

            fields = [
                ("project_type", "项目类型"),
                ("project_period", "项目周期"),
                ("video_duration", "视频时长"),
                ("software_requirements", "软件要求"),
                ("reference_link", "参考链接"),
                ("animation_text_content", "动画文本内容文档"),
                ("voiceover_content", "配音内容"),
            ]

            for field_key, field_label in fields:
                p = doc.add_paragraph()
                p.add_run(f"{field_label}: ").bold = True
                content = str(self.project_info.get(field_key, "未提供"))
                p.add_run(content)

            self.doc_path = f"./client/consultant/{self.project_name}项目文档.docx"
            doc.save(self.doc_path)
            print(f"文档已保存：{self.doc_path}")
        except Exception as e:
            print(f"生成文档时发生错误：{e}")



    # 发送消息
    # def handle_customer_feedback(self, message: str):

if __name__ == "__main__":
    test = ConsultantAgentClient(project_name="反洗钱")
    test.fetch_project_info()
import sqlite3
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt
from langchain_core.tools import tool
from langchain_core.runnables import ensure_config

@tool
def update_project(comments):
    """
    更新项目文档中的信息。

    参数：
        comments (str or list[str]): 客户的意见，可以是字符串或字符串列表。
    """
    config = ensure_config()
    configuration = config.get("configurable", {})
    project_id = configuration.get("project_id", None)
    database = configuration.get("database", None)
    try:
        with sqlite3.connect(database) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT project_doc FROM Projects WHERE project_id = ?", (project_id,))
            result = cursor.fetchone()
            if not result:
                return "未找到对应的项目文档路径。"
            doc_path = result[0]
    except sqlite3.Error as e:
        return f"数据库出现错误: {e}"

    if isinstance(comments, str):
        comments = [comments]

    doc = Document(doc_path)

    update_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    update_paragraph = doc.add_paragraph(update_time, style='Heading 2')
    update_paragraph.add_run(" (请在对应复选框标记完成状态：'☑' 已完成，'☐' 未完成)")

    # 添加每条更新意见和完成状态
    for idx, comment in enumerate(comments, start=1):
        # 默认使用未完成的复选框
        checkbox = '□'

        # 构建文本行
        para = doc.add_paragraph()
        run = para.add_run(checkbox)
        run.font.size = Pt(14)  # 设置复选框字体大小
        para.add_run(f" {idx}. {comment}")  # 添加任务描述

    # 保存文件
    doc.save(doc_path)
    return "客户更新成功。"

@tool
def cancel_project(
) -> str:
    """
    标记项目为已完成。

    返回:
        str: 显示项目是否成功标记为已完成的信息。
    """
    config = ensure_config()
    configuration = config.get("configurable", {})
    project_id = configuration.get("project_id", None)
    database = configuration.get("database", None)
    try:
        with sqlite3.connect(database) as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE Projects SET completion_status = ? WHERE project_id = ?", ("已完成", project_id,))
            conn.commit()
            return "项目标记为已完成"
    except sqlite3.Error as e:
        return "项目标记失败"

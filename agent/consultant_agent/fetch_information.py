import sqlite3
from docx import Document
from langchain_core.tools import tool
from langchain_core.runnables import ensure_config

db = "./db/test.db"

# @tool
def fetch_project_info():
    """
    获取用户的项目需求文档的信息。

    返回：
        一个字符串，其中包含项目文档详细信息和相关的修改意见。
    """
    config = ensure_config()    # Fetch from the context
    configuration = config.get("configurable", {})
    project_id = configuration.get("project_id", None)
    database = configuration.get("database", None)

    with sqlite3.connect(database) as conn:
        cursor = conn.cursor()
        query = """
        SELECT project_doc
        FROM Projects
        WHERE project_id = ?
        """
        cursor.execute(query, (project_id,))
        results = cursor.fetchone()

        doc = Document(results[0])
        docnotes = ""
        for paragraph in doc.paragraphs:
            docnotes += paragraph.text + "\n"

    return docnotes

@tool
def fetch_customer_project_information():   # 一个用户可能会对应多个项目
    """
    获取用户的项目需求文档的信息。

    返回：
        一个字符串，其中包含项目文档详细信息和相关的修改意见。
    """
    config = ensure_config()    # Fetch from the context
    configuration = config.get("configurable", {})
    project_id = configuration.get("project_id", None)
    database = configuration.get("database", None)

    with sqlite3.connect(database) as conn:
        cursor = conn.cursor()
        query = """
        SELECT project_doc
        FROM Projects
        WHERE project_id = ?
        """
        cursor.execute(query, (project_id,))
        results = cursor.fetchone()

        doc = Document(results[0])
        docnotes = ""
        for paragraph in doc.paragraphs:
            docnotes += paragraph.text + "\n"

        cursor.close()
        conn.close()

    return docnotes


@tool
def fetch_supervisor_designer_information():
    """
    获取设计师的项目需求文档的信息。

    返回：
        一个字典列表，其中每个字典包含项目文档详细信息。
    """
    config = ensure_config()
    configuration = config.get("configurable", {})
    designer_id = configuration.get("designer_id", None)
    project_id = configuration.get("project_id", None)
    if not designer_id:
        raise ValueError("No designer ID configured.")
    if not project_id:
        raise ValueError("No project ID configured.")

    conn = sqlite3.connect(db)
    cursor = conn.cursor()

    query = """
    SELECT
        pd.project_type, pd.project_cycle, pd.video_duration, pd.software_use, pd.reference_case, pd.animate_text, pd.voice_content
    FROM
        project_doc pd
    WHERE
        pd.project_id = ?
    """
    cursor.execute(query, (project_id,))
    results = cursor.fetchone()
    results = dict(zip([column[0] for column in cursor.description], results))

    cursor.close()
    conn.close()

    if results is not None:
        return results
    else:
        return f"未找到项目 {project_id} 。"

@tool
def fetch_designer_project_information():
    """
    获取设计师的项目需求文档的信息。

    返回：
        一个字典列表，其中每个字典包含项目文档详细信息。
    """
    config = ensure_config()
    configuration = config.get("configurable", {})
    designer_id = configuration.get("designer_id", None)
    project_id = configuration.get("project_id", None)
    if not designer_id:
        raise ValueError("No designer ID configured.")
    if not project_id:
        raise ValueError("No project ID configured.")

    conn = sqlite3.connect(db)
    cursor = conn.cursor()

    query = """
    SELECT
        pd.project_type, pd.project_cycle, pd.video_duration, pd.software_used, pd.reference_case, pd.animate_text, pd.voice_content
    FROM
        project_doc pd
    WHERE
        pd.project_id = ?
    """
    cursor.execute(query, (project_id,))
    results = cursor.fetchone()
    results = dict(zip([column[0] for column in cursor.description], results))

    cursor.close()
    conn.close()

    if results is not None:
        return results
    else:
        return f"未找到项目 {project_id} 。"